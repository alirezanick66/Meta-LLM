from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from backend.app.utils.custom_normalizer import persian_normalizer
from backend.app.utils.logging_config import LG, LogLevel, log_message
from docx.document import Document as DocxDocument


class WordExtractor:
    """
   ‫ استخراج محتوا از فایل‌های Word (.docx) 
    """

    def __init__( self ):
        self.normalizer = persian_normalizer

    def extract_from_word( self, file_path: str ) -> Tuple[ str, Dict ]:
        """
       ‫ استخراج متن و metadata از فایل Word

        Args:
            file_path: مسیر فایل .docx

        Returns:
            (متن پردازش‌شده و نرمال‌شده, metadata)
        """
        try:
            path = Path( file_path )

            if not path.exists():
                raise FileNotFoundError( f"فایل پیدا نشد: {file_path}" )

            if path.suffix.lower() not in [ '.docx' ]:
                raise ValueError( f"فرمت فایل باید .docx باشد: {file_path}" )

            log_message( LG.DataProcessing, f"شروع استخراج از Word: {path.name}", LogLevel.INFO )

            doc = Document( file_path )
            content_parts: List[ str ] = []

            #‫ ۱. پردازش body اصلی سند (پاراگراف‌ها و جداول به ترتیب)
            self._extract_body( doc, content_parts )

            #‫ ۲. استخراج Textbox ها
            self._extract_textboxes( doc, content_parts )

            # ترکیب همه بخش‌ها
            raw_text = "\n\n".join( part for part in content_parts if part.strip() )

            if not raw_text.strip():
                raise ValueError( "متن استخراج‌شده خالی است" )

            # نرمال‌سازی
            normalized_text = self.normalizer.normalize( raw_text )

            # ساخت metadata
            metadata = self._extract_metadata( path, doc, normalized_text )

            log_message( LG.DataProcessing, f"✅ استخراج Word تکمیل شد - {len(normalized_text)} کاراکتر", LogLevel.INFO )

            return normalized_text, metadata

        except Exception as e:
            log_message( LG.DataProcessing, f"❌ خطا در استخراج Word از {file_path}: {str(e)}", LogLevel.ERROR )
            raise

    def _extract_body( self, doc: DocxDocument, content_parts: List[ str ] ) -> None:
        """
       ‫ پردازش body سند به ترتیب — پاراگراف‌ها و جداول
       ‫ استفاده از qn باعث می‌شود مقایسه تگ‌ها دقیق و سریع‌تر باشد.
        """
        for element in doc.element.body:
            if element.tag == qn( 'w:p' ):
                # پاراگراف معمولی
                para = Paragraph( element, doc )
                text = self._process_paragraph( para )
                if text:
                    content_parts.append( text )

            elif element.tag == qn( 'w:tbl' ):
                # جدول
                table = Table( element, doc )
                table_text = self._process_table( table )
                if table_text:
                    content_parts.append( table_text )

    def _process_paragraph( self, para: Paragraph ) -> str:
        """
        پردازش یک پاراگراف
       ‫ تبدیل Heading به فرمت Markdown
        """
        text = para.text.strip()
        if not text:
            return ""

        # بررسی استایل پاراگراف
        style_name = para.style.name.lower() if para.style and para.style.name else ""

        if "heading 1" in style_name:
            return f"# {text}"
        if "heading 2" in style_name:
            return f"## {text}"
        if "heading 3" in style_name:
            return f"### {text}"
        if "heading 4" in style_name:
            return f"#### {text}"
        if "list" in style_name or style_name.startswith( "list" ):
            return f"- {text}"

        return text

    def _process_table( self, table: Table ) -> str:
        """
        تبدیل جدول به متن ساختاریافته
        بهینه‌سازی شده با استفاده از‫ List Comprehension
        """
        try:
            rows_data: List[ List[ str ] ] = []

            for row in table.rows:
                #‫ استخراج متن سلول‌ها با List Comprehension (سریع‌تر)
                row_cells = [ " ".join( p.text.strip() for p in cell.paragraphs if p.text.strip() ) for cell in row.cells ]
                # فقط ردیف‌هایی که حداقل یک سلول غیرخالی دارند
                if any( row_cells ):
                    rows_data.append( row_cells )

            if not rows_data:
                return ""

            lines = [ "[جدول]" ]

            # فرض: اولین ردیف هدر است
            headers = rows_data[ 0 ]
            data_rows = rows_data[ 1: ]

            if data_rows:
                for row in data_rows:
                    parts = []
                    for i, cell in enumerate( row ):
                        if cell:
                            header = headers[ i ] if i < len( headers ) else f"ستون {i+1}"
                            parts.append( f"{header}: {cell}" )
                    if parts:
                        lines.append( " | ".join( parts ) )
            else:
                # اگر فقط هدر داشتیم
                lines.append( " | ".join( c for c in headers if c ) )

            lines.append( "[/جدول]" )
            return "\n".join( lines )

        except Exception as e:
            log_message( LG.DataProcessing, f"⚠️ خطا در پردازش جدول: {str(e)}", LogLevel.WARNING )
            return ""

    def _extract_textboxes( self, doc: DocxDocument, content_parts: List[ str ] ) -> None:
        """
       ‫ استخراج متن از Textbox ها با استفاده از XPath صحیح
       ‫ این روش بسیار پایدارتر از دستکاری رشته‌ای namespace است.
        """
        try:
            # ‫جستجوی همه txbxContent با استفاده از nsmap
            textbox_contents = doc.element.body.findall( f'.//{qn("w:txbxContent")}' )

            textbox_texts = []
            for tb_content in textbox_contents:
                # ‫پیدا کردن تمام پاراگراف‌های داخل textbox
                for p_elem in tb_content.findall( f'.//{qn("w:p")}' ):
                    para = Paragraph( p_elem, doc )
                    text = para.text.strip()
                    if text:
                        textbox_texts.append( text )

            if textbox_texts:
                combined = "\n".join( textbox_texts )
                content_parts.append( f"[نکته مهم]\n{combined}\n[/نکته مهم]" )
                log_message( LG.DataProcessing, f"✅ {len(textbox_texts)} Textbox استخراج شد", LogLevel.DEBUG )

        except Exception as e:
            #‫ Textbox استخراج نشد ولی این fatal نیست
            log_message( LG.DataProcessing, f"⚠️ خطا در استخراج Textbox: {str(e)}", LogLevel.WARNING )

    def _extract_metadata( self, path: Path, doc: DocxDocument, text: str ) -> Dict:
        """ ‫استخراج metadata از فایل Word"""
        #‫ شمارش هدینگ‌ها با Generator Expression (حفظ حافظه)
        heading_count = sum( 1 for para in doc.paragraphs if para.style and para.style.name and "heading" in para.style.name.lower() )

        table_count = len( doc.tables )

        # ‫عنوان سند — اولین Heading 1 یا اسم فایل
        title = path.stem
        for para in doc.paragraphs:
            if para.style and para.style.name and "heading 1" in para.style.name.lower() and para.text.strip():
                title = para.text.strip()
                break

        return {
            "file_name": path.name,
            "file_path": str( path ),
            "file_type": "docx",
            "title": title,
            "heading_count": heading_count,
            "table_count": table_count,
            "has_tables": table_count > 0,
            "char_count": len( text ),
        }


# Instance سراسری
word_extractor = WordExtractor()
