from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass
import zipfile

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from backend.app.utils.logging_config import log_message, LG, LogLevel
from backend.app.utils.custom_normalizer import persian_normalizer


@dataclass
class DocumentMetadata:
    """متادیتای سند"""
    filename: str
    file_size: int
    total_paragraphs: int
    total_tables: int
    total_headings: int
    has_structure: bool


class TextExtractor:
    """
    استخراج پیشرفته متن از فایل‌های Word با حفظ ساختار
    
    ویژگی‌ها:
    - حفظ headings با hierarchy
    - حفظ bullet points و numbered lists
    - استخراج جداول به markdown
    - بهینه‌سازی performance
    """

    # نگاشت سطح heading به markdown
    HEADING_MAP = {
        'Heading 1': '# ',
        'Heading 2': '## ',
        'Heading 3': '### ',
        'Heading 4': '#### ',
        'Heading 5': '##### ',
        'Heading 6': '###### ',
    }

    def __init__( self ):
        """مقداردهی اولیه"""
        self._current_section = None
        self._section_hierarchy = []
        self._book_title_found = False          # برای تشخیص اولین عنوان به عنوان عنوان کتاب

    def validate_docx( self, file_path: str ) -> bool:
        """
        بررسی معتبر بودن فایل .docx (بهینه شده)
        
        بهبودها:
        - استفاده از zipfile برای validation سریع
        - چک کردن ساختار اصلی بدون load کامل سند
        
        Args:
            file_path: مسیر فایل
            
        Returns:
            True اگر فایل معتبر باشد
        """
        try:
            path = Path( file_path )

            # چک وجود فایل
            if not path.exists():
                log_message( LG.DataProcessing, f"فایل وجود ندارد: {file_path}", LogLevel.ERROR )
                return False

            # چک پسوند
            if path.suffix.lower() != ".docx":
                log_message( LG.DataProcessing, f"فرمت فایل باید .docx باشد: {file_path}", LogLevel.ERROR )
                return False

            # چک ساختار ZIP (فایل‌های docx در واقع ZIP هستند)
            # این روش خیلی سریع‌تر از load کامل سند است
            if not zipfile.is_zipfile( file_path ):
                log_message( LG.DataProcessing, f"فایل DOCX معتبر نیست (ساختار ZIP): {file_path}", LogLevel.ERROR )
                return False

            # چک وجود فایل document.xml (بدون extract کامل)
            with zipfile.ZipFile( file_path, 'r' ) as zip_file:
                if 'word/document.xml' not in zip_file.namelist():
                    log_message( LG.DataProcessing, f"ساختار داخلی DOCX معتبر نیست: {file_path}", LogLevel.ERROR )
                    return False

            log_message( LG.DataProcessing, f"✅ فایل معتبر است: {path.name}", LogLevel.DEBUG )
            return True

        except zipfile.BadZipFile:
            log_message( LG.DataProcessing, f"فایل ZIP خراب است: {file_path}", LogLevel.ERROR )
            return False
        except Exception as e:
            log_message( LG.DataProcessing, f"خطا در اعتبارسنجی {file_path}: {str(e)}", LogLevel.ERROR )
            return False

    def get_file_info( self, file_path: str ) -> Optional[ DocumentMetadata ]:
        """
        دریافت اطلاعات فایل (بهینه شده)
        
        بهبودها:
        - استفاده از properties داخلی docx
        - محاسبه metadata در یک پاس
        - کش کردن نتایج در صورت امکان
        
        Args:
            file_path: مسیر فایل
            
        Returns:
            DocumentMetadata یا None
        """
        try:
            if not self.validate_docx( file_path ):
                return None

            path = Path( file_path )
            doc = DocxDocument( file_path )

            # شمارش در یک پاس
            paragraphs_count = len( doc.paragraphs )
            tables_count = len( doc.tables )

            # شمارش headings

            headings_count = sum( 1 for para in doc.paragraphs
                                  if para.style.name.startswith( 'Heading' ) )          #type: ignore

            metadata = DocumentMetadata( filename=path.name,
                                         file_size=path.stat().st_size,
                                         total_paragraphs=paragraphs_count,
                                         total_tables=tables_count,
                                         total_headings=headings_count,
                                         has_structure=headings_count > 0 or tables_count > 0 )

            log_message(
                LG.DataProcessing, f"📄 {metadata.filename}: {metadata.total_paragraphs} پاراگراف، "
                f"{metadata.total_tables} جدول، {metadata.total_headings} عنوان", LogLevel.INFO )

            return metadata

        except Exception as e:
            log_message( LG.DataProcessing, f"خطا در دریافت اطلاعات فایل: {str(e)}", LogLevel.ERROR )
            return None

    def _extract_table( self, table: Table ) -> str:
        """
        تبدیل جدول به فرمت markdown
        
        Args:
            table: جدول docx
            
        Returns:
            جدول به فرمت markdown
        """
        try:
            if not table.rows:
                return ""

            lines = []

            # استخراج header (اولین ردیف)
            header_cells = [ cell.text.strip() for cell in table.rows[ 0 ].cells ]
            lines.append( "| " + " | ".join( header_cells ) + " |" )

            # خط جدا کننده
            lines.append( "| " + " | ".join( [ "---" ] * len( header_cells ) ) + " |" )

            # بقیه ردیف‌ها
            for row in table.rows[ 1: ]:
                cells = [ cell.text.strip() for cell in row.cells ]
                lines.append( "| " + " | ".join( cells ) + " |" )

            return "\n".join( lines )

        except Exception as e:
            log_message( LG.DataProcessing, f"خطا در استخراج جدول: {str(e)}", LogLevel.WARNING )
            return "[جدول قابل استخراج نیست]"

    def _is_list_paragraph( self, paragraph: Paragraph ) -> Tuple[ bool, str ]:
        """
        تشخیص list items (bullet یا numbered)
        
        Args:
            paragraph: پاراگراف docx
            
        Returns:
            (is_list, prefix) - آیا لیست هست و پیشوند آن
        """
        try:
            # چک کردن numbering
            if paragraph._element.pPr is not None:
                numPr = paragraph._element.pPr.numPr
                if numPr is not None:
                    # اگه numPr داشته باشه، یعنی لیست است
                    numId = numPr.numId
                    if numId is not None:
                        # تشخیص نوع: bullet یا numbered
                        # برای سادگی از bullet استفاده می‌کنیم
                        return True, "• "

            # fallback: چک کردن متن
            text = paragraph.text.strip()
            if text and len( text ) > 0:
                first_char = text[ 0 ]
                # چک کردن bullet های معمول
                if first_char in [ '•', '·', '○', '■', '□', '-', '*' ]:
                    return True, "• "

            return False, ""

        except Exception:
            return False, ""

    def _detect_heading_type( self, text: str, style_name: str ) -> str:
        """
        تشخیص نوع heading بر اساس محتوا (صرف‌نظر از style level)
        
        Args:
            text: متن heading
            style_name: نام style از Word
            
        Returns:
            نوع heading (book_title, author, chapter, subsection)
        """
        text_lower = text.lower().strip()

        # 1. تشخیص نویسنده (اولویت اول)
        author_patterns = [ 'نویسنده', 'تألیف', 'تاليف', 'مؤلف', 'مولف', 'نوشته', 'از:', 'author' ]
        if any( pattern in text_lower for pattern in author_patterns ):
            return 'author'

        # 2. تشخیص فصل (بر اساس محتوا)
        chapter_patterns = [ 'فصل', 'بخش', 'قسمت', 'باب', 'chapter', 'part' ]
        if any( pattern in text_lower for pattern in chapter_patterns ):
            return 'chapter'

        # 3. اولین heading که دیدیم = عنوان کتاب (صرف‌نظر از style)
        if not self._book_title_found:
            self._book_title_found = True
            return 'book_title'

        # 4. بقیه headingها = زیربخش
        return 'subsection'

    def _get_heading_level( self, paragraph: Paragraph ) -> Optional[ str ]:
        """
        تشخیص سطح heading با لاجیک هوشمند (بر اساس محتوا)
        
        Args:
            paragraph: پاراگراف docx
            
        Returns:
            markdown prefix یا None
        """
        style_name = paragraph.style.name

        # چک کردن اینکه آیا این یک heading است
        if not style_name.startswith( 'Heading' ):
            return self.HEADING_MAP.get( style_name )

        # تشخیص نوع heading (بر اساس محتوا، نه style)
        heading_type = self._detect_heading_type( paragraph.text, style_name )

        # نگاشت بر اساس نوع محتوا
        if heading_type == 'book_title':
            return '# '          # عنوان کتاب: #
        elif heading_type == 'author':
            return '## '          # نام نویسنده: ##
        elif heading_type == 'chapter':
            return '## '          # عنوان فصل: ##
        elif heading_type == 'subsection':
            return '### '          # زیربخش‌ها: ###

        # fallback به نگاشت پیش‌فرض
        return self.HEADING_MAP.get( style_name )

    def _is_chapter_paragraph( self, text: str ) -> bool:
        """
        تشخیص پاراگراف‌های معمولی که فصل هستند (بدون style)
        
        Args:
            text: متن پاراگراف
            
        Returns:
            True اگر این پاراگراف یک فصل باشد
        """
        text_stripped = text.strip()

        # چک کردن الگوهای فصل
        chapter_patterns = [
            r'^فصل\s+(اول|دوم|سوم|چهارم|پنجم|ششم|هفتم|هشتم|نهم|دهم)',
            r'^فصل\s+\d+',
            r'^فصل\s+[۰-۹]+',
            r'^chapter\s+\d+',
        ]

        import re
        for pattern in chapter_patterns:
            if re.match( pattern, text_stripped, re.IGNORECASE ):
                return True

        return False

    def extract_from_docx( self, file_path: str ) -> Tuple[ str, DocumentMetadata ]:
        """
        استخراج متن با حفظ ساختار کامل
        
        Args:
            file_path: مسیر فایل
            
        Returns:
            (متن استخراج شده, metadata)
        """
        try:
            # دریافت metadata
            metadata = self.get_file_info( file_path )
            if not metadata:
                raise ValueError( "فایل معتبر نیست" )

            doc = DocxDocument( file_path )
            extracted_parts = []

            # ریست کردن state برای سند جدید
            self._book_title_found = False

            # ساخت dictionary سریع برای دسترسی به paragraph objects
            para_dict = { p._element: p for p in doc.paragraphs }
            table_dict = { t._element: t for t in doc.tables }

            # پردازش هر element به ترتیب
            for element in doc.element.body:
                # پاراگراف
                if element.tag.endswith( 'p' ):
                    para = para_dict.get( element )

                    if para and para.text.strip():
                        # چک heading
                        heading_prefix = self._get_heading_level( para )
                        if heading_prefix:
                            extracted_parts.append( f"\n{heading_prefix}{para.text.strip()}\n" )
                            continue

                        # چک list
                        is_list, list_prefix = self._is_list_paragraph( para )
                        if is_list:
                            extracted_parts.append( f"{list_prefix}{para.text.strip()}" )
                            continue

                        # چک فصل بدون style
                        if self._is_chapter_paragraph( para.text ):
                            extracted_parts.append( f"\n## {para.text.strip()}\n" )
                            continue

                        # پاراگراف معمولی
                        extracted_parts.append( para.text.strip() )

                # جدول
                elif element.tag.endswith( 'tbl' ):
                    table = table_dict.get( element )

                    if table:
                        table_md = self._extract_table( table )
                        if table_md:
                            extracted_parts.append( f"\n{table_md}\n" )

            # ترکیب نهایی
            full_text = "\n".join( extracted_parts )

            # پاکسازی خطوط خالی اضافی
            full_text = "\n".join( line for line in full_text.split( "\n" )
                                   if line.strip() or line == ""          # حفظ یک خط خالی
                                  )

            # نرمال‌سازی متن فارسی
            full_text = persian_normalizer.normalize( full_text )

            log_message( LG.DataProcessing,
                         f"✅ استخراج و نرمال‌سازی کامل شد: {len(full_text)} کاراکتر از {metadata.filename}",
                         LogLevel.INFO )

            return full_text, metadata

        except Exception as e:
            log_message( LG.DataProcessing, f"خطا در استخراج متن از {file_path}: {str(e)}", LogLevel.ERROR )
            raise


# Instance سراسری
text_extractor = TextExtractor()
